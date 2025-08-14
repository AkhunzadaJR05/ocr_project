import streamlit as st
import pytesseract
from PIL import Image
from docx import Document
import shutil
import re
import os
from io import BytesIO

TEMPLATE_PATH = "sales_receipt.docx"

# Predefined lists for more robust Make and Model extraction
COMMON_MAKES = {
    "VOLKSWAGEN": ["VOLKSWAGEN", "VW"],
    "FORD": ["FORD"],
    "TOYOTA": ["TOYOTA"],
    "HONDA": ["HONDA"],
    "BMW": ["BMW"],
    "MERCEDES-BENZ": ["MERCEDES", "BENZ"],
}

COMMON_MODELS = {
    "GOLF": ["GOLF"],
    "FIESTA": ["FIESTA"],
    "FOCUS": ["FOCUS"],
    "CIVIC": ["CIVIC"],
    "COROLLA": ["COROLLA"]
}

def safe_extract(text, pattern, group=1):
    """
    Extracts data using a regex pattern, returning None if no match is found.
    """
    try:
        match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        if match:
            return match.group(group).strip()
        return None
    except Exception:
        return None

# Function to extract the registration number
def extract_reg_number(text):
    """
    Finds and extracts a UK registration number using multiple patterns.
    """
    # More flexible regex patterns for UK number plates
    reg_number_patterns = [
        r"[A-Z]{2}\s*\d{2}\s*[A-Z]{3}",  # Common format (e.g., GD65 EGF)
        r"Registration\s*Number\s*([A-Z0-9\s]+)", # Looks for the "Registration Number" label
        r"([A-Z]{2}\d{2}[A-Z]{3})", # Handles no-space format from OCR errors
    ]
    
    for pattern in reg_number_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            # Clean and return the extracted value
            return re.sub(r"[^A-Z0-9]", "", match.group(0).upper())

    # Fallback to a position-based search below the "Registration Number" label
    chunks = text.split("Registration Number")
    if len(chunks) > 1:
        next_lines = chunks[1].splitlines()[:5]
        for line in next_lines:
            clean_line = re.sub(r"[^A-Z0-9]", "", line.upper())
            if len(clean_line) >= 7:
                return clean_line[:7]

    return "N/A"

# Main function for extracting all vehicle data
def extract_data_from_image(image):
    """
    Orchestrates the OCR process and data extraction from the uploaded image.
    """
    tesseract_path = shutil.which('tesseract')
    if tesseract_path:
        pytesseract.pytesseract.tesseract_cmd = tesseract_path
    else:
        st.error("Tesseract not found. Please check your packages.txt file.")
        st.stop()

    text = pytesseract.image_to_string(image)
    
    # Use a more robust chain of patterns for each field
    data = {
        "make": "N/A",
        "model": "N/A",
        # New, more flexible patterns for 'year'
        "year": safe_extract(text, r"(?:Date of first|B: Date).*?(\d{4})") or 
                safe_extract(text, r"B\s*:\s*Date of first.*?(\d{4})") or 
                safe_extract(text, r"(\d{4})")
    }

    # New logic to extract make from predefined list
    for make, keywords in COMMON_MAKES.items():
        for keyword in keywords:
            if re.search(r'\b' + re.escape(keyword) + r'\b', text, re.IGNORECASE):
                data["make"] = make
                break
        if data["make"] != "N/A":
            break

    # New logic to extract model from predefined list
    for model, keywords in COMMON_MODELS.items():
        for keyword in keywords:
            if re.search(r'\b' + re.escape(keyword) + r'\b', text, re.IGNORECASE):
                data["model"] = model
                break
        if data["model"] != "N/A":
            break


    # Backup extraction for the Chassis number, which is very specific
    chasis_raw = safe_extract(text, r"E\s*:?\s*VIN.*([A-Z0-9]{17})")
    data["chasis"] = chasis_raw
    
    # Post-processing for chassis number
    if data["chasis"]:
        # Check for the specific VW OCR error (WVW is misread as VW2)
        if len(data["chasis"]) == 17 and data["chasis"].startswith("VW2"):
            data["chasis"] = "WVW" + data["chasis"][3:]

    # Extract the registration number
    data["reg_number"] = extract_reg_number(text)

    # Post-processing to clean up extracted data and assign 'N/A' as a default
    if not data["year"]:
        data["year"] = "N/A"
    
    return data

# Function to fill a Word document template
def fill_word_template(data):
    """
    Fills a Word document template with the extracted data.
    """
    doc = Document(TEMPLATE_PATH)
    
    replacements = {
        "{{make}}": data["make"],
        "{{model}}": data["model"],
        "{{year}}": data["year"],
        "{{chasis}}": data["chasis"],
        "{{reg_number}}": data["reg_number"]
    }

    # Replace placeholders in paragraphs
    for p in doc.paragraphs:
        for key, value in replacements.items():
            if key in p.text:
                p.text = p.text.replace(key, value)

    # Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)

    # Save the filled document to a buffer and return it
    output_buffer = BytesIO()
    doc.save(output_buffer)
    output_buffer.seek(0)
    return output_buffer

# --- Streamlit UI ---
st.title("Vehicle Receipt OCR Generator")
uploaded_image = st.file_uploader("Upload vehicle document image", type=["jpg", "jpeg", "png"])

if uploaded_image is not None:
    img = Image.open(uploaded_image)
    st.image(img, caption="Uploaded Image", use_container_width=True)

    with st.spinner("Extracting data..."):
        extracted_data = extract_data_from_image(img)

    st.subheader("Extracted Data")
    st.json(extracted_data)

    if st.button("Generate Receipt"):
        receipt_file = fill_word_template(extracted_data)
        st.success("Receipt generated successfully!")
        st.download_button(
            label="Download Word Receipt",
            data=receipt_file,
            file_name="filled_receipt.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
